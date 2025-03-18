import os
import spacy
import pdfplumber
import docx
import re
from django.shortcuts import render, redirect
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from pymongo import MongoClient
from django.http import JsonResponse
from bson.objectid import ObjectId


# Connect to MongoDB
client = MongoClient("mongodb://localhost:27017/")  # Update URI if needed
resume_db = client["json"]  # Resume Database
job_db = client["kaggle_db"]  # Job Database

resume_collection = resume_db["ParsedResumes"]  # Resumes Collection
job_collection = job_db["job_postings"]  # Job Postings Collection
matched_collection = resume_db["MatchedJobs"]  # Collection to store matched jobs

# Load spaCy NLP model
nlp = spacy.load("en_core_web_sm")

# Define keywords for parsing
SKILL_KEYWORDS = {"Python", "Java", "C++", "Flask", "Django", "Machine Learning", "Deep Learning", "SQL", "NoSQL", "AWS", "Terraform", "Linux","SQL", "Azure", "Data Factory", "DBT", "ETL", "Data structures", "Algorithms", "Async programming", "Object-oriented design", "Parallel programming", ".NET", "Java", "Angular", "Git", "Cloud-based systems", "Big data design", "Data normalization", "Queuing technologies", "Metrics", "Logging", "Monitoring", "Alerting", "RESTful APIs", "HL7 V2.x / FHIR", "CI/CD Pipeline", "K8s", "Terraform", "Electronic Health Records"
}
EXPERIENCE_KEYWORDS = {"intern", "developer", "engineer", "full-time", "research"}
PROJECT_KEYWORDS = {"project", "developed", "implemented", "designed"}

# Threshold for job matching
THRESHOLD = 50 # Percentage match


def index(request):
    """Render the upload page."""
    return render(request, 'index1.html')

def page2(request):
    matched_jobs = request.session.get('matched_jobs', [])
    return render(request, 'page2.html', {'matched_jobs': matched_jobs})
    


def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + "\n"
    return text


def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])


def extract_resume_details(text):
    """Extract skills, projects, and experience from resume text."""
    skills, projects, experience = set(), set(), set()
    
    lines = text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        for skill in SKILL_KEYWORDS:
            if re.search(rf"\b{re.escape(skill)}\b", line, re.IGNORECASE):
                skills.add(skill)

        if any(word in line.lower() for word in PROJECT_KEYWORDS):
            projects.add(line)

        if any(word in line.lower() for word in EXPERIENCE_KEYWORDS):
            experience.add(line)

    return {
        "Projects": list(projects),
        "Experience": list(experience),
        "Skills": sorted(skills)
    }


def store_resume_in_mongodb(parsed_data):
    """Store parsed resume data in MongoDB."""
    result = resume_collection.insert_one(parsed_data)
    print(f"Resume data inserted with ID: {result.inserted_id}")
    return str(result.inserted_id)  # Return the MongoDB document ID


def calculate_match(resume_skills, job_skills):
    """Calculate skill match percentage between resume and job posting."""
    if not isinstance(job_skills, str):  
        job_skills = ""

    resume_set = set(resume_skills)
    job_set = set(job_skills.split(", "))  # Assuming skills are comma-separated

    if not job_set:
        return 0

    common_skills = resume_set.intersection(job_set)
    match_percentage = (len(common_skills) / len(job_set)) * 100
    return match_percentage


def match_resume_to_jobs(resume_id, resume_skills):
    """Find and store job postings that match the resume skills."""
    matched_jobs = []

    jobs = list(job_collection.find())  # Fetch all job postings
    for job in jobs:
        job_title = job.get("job_title", "Unknown Job")
        job_company = job.get("company", "Unknown Company")
        job_location = job.get("job_location", "Unknown Location")
        job_link = job.get("job_link", "#")
        job_skills = job.get("job_skills", "")

        match_score = calculate_match(resume_skills, job_skills)

        if match_score >= THRESHOLD:
            matched_jobs.append({
                "job_title": job_title,
                "company": job_company,
                "location": job_location,
                "job_link": job_link,
                "match_score": round(match_score, 2)
            })

    if matched_jobs:
        matched_collection.insert_one({
            "resume_id": resume_id,
            "matched_jobs": matched_jobs
        })
        print(f"Matched jobs stored for Resume ID: {resume_id}")

    return matched_jobs


def upload_and_analyze(request):
    """Handle resume upload, parse details, store in DB, and match with jobs."""
    if request.method == "POST" and request.FILES.get("resume"):
        uploaded_file = request.FILES["resume"]
        file_path = default_storage.save(f"uploads/{uploaded_file.name}", ContentFile(uploaded_file.read()))
        full_path = os.path.join(default_storage.location, file_path)

        # Extract text based on file type
        if uploaded_file.name.endswith(".pdf"):
            resume_text = extract_text_from_pdf(full_path)
        elif uploaded_file.name.endswith(".docx"):
            resume_text = extract_text_from_docx(full_path)
        else:
            return JsonResponse({"error": "Unsupported file format"}, status=400)

        # Parse resume details
        parsed_data = extract_resume_details(resume_text)
        resume_id = store_resume_in_mongodb(parsed_data)  # Store in DB and get ID

        # Match with jobs
        matched_jobs = match_resume_to_jobs(resume_id, parsed_data["Skills"])

        # Store matched jobs in session to pass to page2
        request.session["matched_jobs"] = matched_jobs
        
        # ✅ Redirect to page2 after processing
        return redirect('page2')

    return JsonResponse({"error": "No file uploaded"}, status=400)

def get_matched_jobs(request):
    """Fetch only the latest matched jobs from MongoDB."""
    latest_match = matched_collection.find_one({}, {"_id": 0, "resume_id": 1, "matched_jobs": 1}, sort=[("_id", -1)])

    if not latest_match:
        return JsonResponse({"results": []}, safe=False)

    resume_id = latest_match.get("resume_id", "Unknown ID")
    jobs = latest_match.get("matched_jobs", [])

    job_list = [
        {
            "job_title": job.get("job_title", "Unknown Job"),
            "location": job.get("location", "Unknown Location"),
            "match_score": job.get("match_score", 0)
        }
        for job in jobs
    ]

    return JsonResponse({"results": [{"resume_id": resume_id, "matched_jobs": job_list}]}, safe=False)
def upload_resume(request):
    if request.method == 'POST':
        # Handle the uploaded resume file (optional)
        uploaded_file = request.FILES.get('resume')

        # Redirect to page2.html
        return redirect('page2')  # 'page2' should be the name of the URL pattern
    #return render(request, 'index.html')


# import os
# import spacy
# import pdfplumber
# import docx
# import json
# import re
# from django.shortcuts import render, redirect
# from django.core.files.storage import default_storage
# from django.core.files.base import ContentFile
# from pymongo import MongoClient
# from django.http import JsonResponse
# from bson.objectid import ObjectId

# # Connect to MongoDB
# client = MongoClient("mongodb://localhost:27017/")  # Update URI if needed
# resume_db = client["json"]  # Resume Database
# job_db = client["kaggle_db"]  # Job Database

# resume_collection = resume_db["ParsedResumes"]  # Resumes Collection
# job_collection = job_db["job_postings"]  # Job Postings Collection
# matched_collection = resume_db["MatchedJobs"]  # Collection to store matched jobs

# # Load spaCy NLP model
# nlp = spacy.load("en_core_web_sm")

# # Define keywords for parsing
# SKILL_KEYWORDS = {
#     "Python", "Java", "C++", "Flask", "Django", "Machine Learning", "Deep Learning", "SQL", "NoSQL", "AWS",
#     "Terraform", "Linux", "Azure", "Data Factory", "DBT", "ETL", "Data Structures", "Algorithms", "Async programming",
#     "Object-oriented design", "Parallel programming", ".NET", "Angular", "Git", "Cloud-based systems",
#     "Big data design", "Data normalization", "Queuing technologies", "Metrics", "Logging", "Monitoring", "Alerting",
#     "RESTful APIs", "HL7 V2.x / FHIR", "CI/CD Pipeline", "K8s", "Electronic Health Records"
# }

# EXPERIENCE_KEYWORDS = {"intern", "internship", "developer", "engineer", "full-time", "research", "fellow", "hackathon"}
# PROJECT_KEYWORDS = {"project", "built", "developed", "implemented", "designed"}

# IGNORE_KEYWORDS = {
#     "achievements", "certifications", "solved", "rating", "hackerrank", "GFG", "award", "university", "college",
#     "cgpa", "gpa", "btech", "bachelor", "engineering", "school", "secondary", "club", "gender", "hobbies", "languages",
#     "reference", "date of birth", "nationality", "phone", "email", "designation", "professor", "head of department",
#     "HOD", "co-curricular", "extra-curricular", "representative", "paper presentation", "nss", "technical club",
#     "non-technical club"
# }

# # Threshold for job matching
# THRESHOLD = 60  # Percentage match


# def index(request):
#     """Render the upload page."""
#     return render(request, 'index1.html')  # Ensure 'index1.html' exists in the templates folder


# def page2(request):
#     """Render the job matches page."""
#     matched_jobs = request.session.get('matched_jobs', [])
#     return render(request, 'page2.html', {'matched_jobs': matched_jobs})


# def extract_text_from_pdf(pdf_path):
#     """Extract text from a PDF file."""
#     text = ""
#     with pdfplumber.open(pdf_path) as pdf:
#         for page in pdf.pages:
#             extracted_text = page.extract_text()
#             if extracted_text:
#                 text += extracted_text + "\n"
#     return text


# def extract_text_from_docx(docx_path):
#     """Extract text from a DOCX file."""
#     doc = docx.Document(docx_path)
#     return "\n".join([para.text for para in doc.paragraphs])


# def extract_resume_details(text):
#     """Extract skills, projects, and experience from resume text."""
#     doc = nlp(text)
#     skills, projects, experience = set(), set(), set()

#     lines = text.split("\n")  # Process text line by line
#     in_experience_section = False
#     in_projects_section = False

#     for line in lines:
#         line = line.strip()
#         if not line:
#             continue

#         # Ignore unwanted sections
#         if any(word in line.lower() for word in IGNORE_KEYWORDS):
#             continue

#         # Detect section headers
#         if "project" in line.lower():
#             in_experience_section = False
#             in_projects_section = True
#             continue
#         elif "experience" in line.lower() or "internship" in line.lower():
#             in_experience_section = True
#             in_projects_section = False
#             continue

#         # Extract skills
#         for skill in SKILL_KEYWORDS:
#             if re.search(rf"\b{re.escape(skill)}\b", line, re.IGNORECASE):
#                 skills.add(skill)

#         # Extract projects
#         if in_projects_section or any(word in line.lower() for word in PROJECT_KEYWORDS):
#             projects.add(line)

#         # Extract experience
#         elif in_experience_section or any(word in line.lower() for word in EXPERIENCE_KEYWORDS):
#             experience.add(line)

#     return {
#         "Projects": list(projects),
#         "Experience": list(experience),
#         "Skills": sorted(skills)
#     }


# def store_resume_in_mongodb(parsed_data):
#     """Store parsed resume data in MongoDB."""
#     result = resume_collection.insert_one(parsed_data)
#     print(f"Resume data inserted with ID: {result.inserted_id}")
#     return str(result.inserted_id)  # Return the MongoDB document ID


# def match_resume_to_jobs(resume_id, resume_skills):
#     """Find and store job postings that match the resume skills."""
#     matched_jobs = []

#     jobs = list(job_collection.find())  # Fetch all job postings
#     for job in jobs:
#         job_title = job.get("job_title", "Unknown Job")
#         job_company = job.get("company", "Unknown Company")
#         job_location = job.get("job_location", "Unknown Location")
#         job_link = job.get("job_link", "#")

#         # Ensure job_skills is a string before using split
#         job_skills = job.get("job_skills", "")
#         if not isinstance(job_skills, str):  # Convert to empty string if it's not a string
#             job_skills = ""

#         job_skills_list = job_skills.split(", ")  # Assuming skills are comma-separated

#         if not job_skills_list:  # Prevent division by zero
#             match_score = 0
#         else:
#             common_skills = set(resume_skills).intersection(job_skills_list)
#             match_score = (len(common_skills) / len(job_skills_list)) * 100

#         if match_score >= THRESHOLD:
#             matched_jobs.append({
#                 "job_title": job_title,
#                 "company": job_company,
#                 "location": job_location,
#                 "job_link": job_link,
#                 "match_score": round(match_score, 2)
#             })

#     if matched_jobs:
#         matched_collection.insert_one({
#             "resume_id": resume_id,
#             "matched_jobs": matched_jobs
#         })
#         print(f"Matched jobs stored for Resume ID: {resume_id}")

#     return matched_jobs



# def upload_and_analyze(request):
#     """Handle resume upload, parse details, store in DB, and match with jobs."""
#     if request.method == "POST" and request.FILES.get("resume"):
#         uploaded_file = request.FILES["resume"]
#         file_path = default_storage.save(f"uploads/{uploaded_file.name}", ContentFile(uploaded_file.read()))
#         full_path = os.path.join(default_storage.location, file_path)

#         # Extract text based on file type
#         if uploaded_file.name.endswith(".pdf"):
#             resume_text = extract_text_from_pdf(full_path)
#         elif uploaded_file.name.endswith(".docx"):
#             resume_text = extract_text_from_docx(full_path)
#         else:
#             return JsonResponse({"error": "Unsupported file format"}, status=400)

#         # Parse resume details
#         parsed_data = extract_resume_details(resume_text)
#         resume_id = store_resume_in_mongodb(parsed_data)

#         # Match with jobs
#         matched_jobs = match_resume_to_jobs(resume_id, parsed_data["Skills"])

#         # Store matched jobs in session
#         request.session["matched_jobs"] = matched_jobs

#         return redirect('page2')

#     return JsonResponse({"error": "No file uploaded"}, status=400)


# def get_matched_jobs(request):
#     """Fetch matched jobs from MongoDB."""
#     latest_match = matched_collection.find_one({}, {"_id": 0, "resume_id": 1, "matched_jobs": 1}, sort=[("_id", -1)])
#     return JsonResponse({"results": [latest_match] if latest_match else []}, safe=False)


# def upload_resume(request):
#     """Redirect to page2 after resume upload."""
#     if request.method == 'POST':
#         return redirect('page2')
